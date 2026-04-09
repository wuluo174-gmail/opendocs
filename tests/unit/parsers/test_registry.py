"""Tests for ParserRegistry and failure isolation."""

from __future__ import annotations

from pathlib import Path

import pytest

from opendocs.indexing.chunker import Chunker
from opendocs.parsers import create_default_registry
from opendocs.parsers.base import Paragraph, ParserRegistry
from opendocs.parsers.txt_parser import TxtParser

TEST_DOC_ID = "22222222-2222-4222-8222-222222222222"


def _chunk(parsed, *, doc_id: str = TEST_DOC_ID):
    return Chunker().chunk(parsed, doc_id=doc_id)


class TestParserRegistry:
    def test_register_and_lookup(self) -> None:
        registry = ParserRegistry()
        registry.register(TxtParser())
        assert registry.is_supported("test.txt")
        assert registry.get_parser("test.txt") is not None

    def test_unsupported_format(self) -> None:
        registry = ParserRegistry()
        assert not registry.is_supported("test.xyz")
        assert registry.get_parser("test.xyz") is None

    def test_create_default_registry(self) -> None:
        registry = create_default_registry()
        assert registry.is_supported("file.txt")
        assert registry.is_supported("file.md")
        assert registry.is_supported("file.docx")
        assert registry.is_supported("file.pdf")
        assert not registry.is_supported("file.xyz")

    def test_doc_format_not_supported(self) -> None:
        """Spec §24: .doc is explicitly not supported, only .docx."""
        registry = create_default_registry()
        assert not registry.is_supported("legacy.doc")
        assert registry.get_parser("legacy.doc") is None


class TestFailureIsolation:
    """Registry.parse() should never raise – it returns failed ParsedDocument."""

    def test_unsupported_format_no_exception(self, tmp_path: Path) -> None:
        registry = create_default_registry()
        p = tmp_path / "test.xyz"
        p.write_text("hello")
        result = registry.parse(p)
        assert result.parse_status == "failed"
        assert result.file_type == "unsupported"
        assert "unsupported format" in result.error_info
        assert result.error is not None
        assert result.error.code == "unsupported_format"
        assert result.error.details["extension"] == ".xyz"

    def test_empty_file_returns_failed(self, tmp_path: Path) -> None:
        """TC-002: empty files should surface in the failure bucket."""
        registry = create_default_registry()
        p = tmp_path / "empty.txt"
        p.write_text("")
        result = registry.parse(p)
        assert result.parse_status == "failed"
        assert result.file_type == "txt"
        assert result.raw_text == ""
        assert result.paragraphs == []
        assert result.error is not None
        assert result.error.code == "no_extractable_text"
        assert result.error.details["parser"] == "TxtParser"

    def test_whitespace_only_txt_returns_failed(self, tmp_path: Path) -> None:
        registry = create_default_registry()
        p = tmp_path / "blank.txt"
        p.write_text(" \n\t \n", encoding="utf-8")

        result = registry.parse(p)

        assert result.parse_status == "failed"
        assert result.raw_text == ""
        assert result.paragraphs == []
        assert result.error is not None
        assert result.error.code == "no_extractable_text"
        assert result.error.details["parser"] == "TxtParser"

    def test_whitespace_only_markdown_returns_failed(self, tmp_path: Path) -> None:
        registry = create_default_registry()
        p = tmp_path / "blank.md"
        p.write_text("  \n\n\t\n", encoding="utf-8")

        result = registry.parse(p)

        assert result.parse_status == "failed"
        assert result.raw_text == ""
        assert result.paragraphs == []
        assert result.error is not None
        assert result.error.code == "no_extractable_text"
        assert result.error.details["parser"] == "MdParser"

    def test_whitespace_only_docx_returns_failed(self, tmp_path: Path) -> None:
        docx = pytest.importorskip("docx")
        Document = docx.Document  # type: ignore[attr-defined]
        registry = create_default_registry()
        doc = Document()
        doc.add_paragraph("   ")
        p = tmp_path / "blank.docx"
        doc.save(str(p))

        result = registry.parse(p)

        assert result.parse_status == "failed"
        assert result.raw_text == ""
        assert result.paragraphs == []
        assert result.error is not None
        assert result.error.code == "no_extractable_text"
        assert result.error.details["parser"] == "DocxParser"

    def test_direct_parser_and_registry_share_same_final_contract(
        self,
        tmp_path: Path,
    ) -> None:
        p = tmp_path / "blank.txt"
        p.write_text("", encoding="utf-8")

        direct = TxtParser().parse(p)
        via_registry = create_default_registry().parse(p)

        assert direct.model_dump() == via_registry.model_dump()

    def test_missing_file_no_exception(self, tmp_path: Path) -> None:
        registry = create_default_registry()
        p = tmp_path / "nonexistent.txt"
        result = registry.parse(p)
        assert result.parse_status == "failed"
        assert result.error_info is not None
        assert result.error is not None
        assert result.error.code == "io_error"

    def test_corrupted_docx_no_exception(self, tmp_path: Path) -> None:
        registry = create_default_registry()
        p = tmp_path / "bad.docx"
        p.write_bytes(b"not a docx file")
        result = registry.parse(p)
        assert result.parse_status == "failed"
        assert result.error_info is not None
        assert result.error is not None
        assert result.error.code == "parse_failed"
        assert result.error.details["parser"] == "DocxParser"

    def test_corrupted_pdf_no_exception(self, tmp_path: Path) -> None:
        registry = create_default_registry()
        p = tmp_path / "bad.pdf"
        p.write_bytes(b"not a pdf")
        result = registry.parse(p)
        assert result.parse_status == "failed"
        assert result.error_info is not None
        assert result.error is not None
        assert result.error.code == "parse_failed"
        assert result.error.details["parser"] == "PdfParser"

    def test_permission_denied_no_exception(self, tmp_path: Path) -> None:
        registry = create_default_registry()
        p = tmp_path / "noperm.txt"
        p.write_text("content")
        p.chmod(0o000)
        try:
            result = registry.parse(p)
            assert result.parse_status == "failed"
            assert result.error is not None
        finally:
            p.chmod(0o644)

    def test_doc_format_rejected(self, tmp_path: Path) -> None:
        """Spec §24: .doc files must be rejected with failed status."""
        registry = create_default_registry()
        p = tmp_path / "legacy.doc"
        p.write_bytes(b"\xd0\xcf\x11\xe0")  # OLE magic bytes
        result = registry.parse(p)
        assert result.parse_status == "failed"
        assert result.file_type == "unsupported"
        assert "unsupported format" in result.error_info
        assert result.error is not None
        assert result.error.code == "unsupported_format"

    def test_batch_processing_never_crashes(self, tmp_path: Path) -> None:
        """Simulate batch: mix of good, bad, and unsupported files."""
        registry = create_default_registry()

        good = tmp_path / "good.txt"
        good.write_text("Hello world")

        bad = tmp_path / "bad.docx"
        bad.write_bytes(b"corrupt")

        unsup = tmp_path / "file.xyz"
        unsup.write_text("data")

        legacy = tmp_path / "old.doc"
        legacy.write_bytes(b"\xd0\xcf\x11\xe0")  # OLE magic bytes

        results = [registry.parse(f) for f in [good, bad, unsup, legacy]]
        # No exception raised – all results returned
        assert len(results) == 4
        assert results[0].parse_status == "success"
        assert results[1].parse_status == "failed"
        assert results[2].parse_status == "failed"
        assert results[3].parse_status == "failed"  # .doc rejected
        assert results[2].file_type == "unsupported"
        assert results[3].file_type == "unsupported"
        assert results[1].error is not None
        assert results[2].error is not None
        assert results[3].error is not None

    def test_normalization_applied(self, tmp_path: Path) -> None:
        """Registry.parse() must apply text normalization."""
        registry = create_default_registry()
        p = tmp_path / "fullwidth.txt"
        # Full-width ABC and full-width space
        p.write_text("ＡＢＣ\u3000标题\n\n第一段。", encoding="utf-8")
        result = registry.parse(p)
        assert result.parse_status == "success"
        # Full-width chars should be normalized to half-width
        assert "ABC" in result.raw_text
        assert "\u3000" not in result.raw_text
        assert result.title == "ABC 标题"
        for para in result.paragraphs:
            assert "Ａ" not in para.text

    def test_heading_path_is_normalized(self, tmp_path: Path) -> None:
        """Heading paths should use the same normalized form as title/raw_text."""
        registry = create_default_registry()
        p = tmp_path / "heading_norm.md"
        p.write_text("# ＡＢＣ  标题\n\n正文内容。", encoding="utf-8")

        result = registry.parse(p)

        assert result.parse_status == "success"
        assert result.title == "ABC 标题"
        assert [para.heading_path for para in result.paragraphs] == ["ABC 标题", "ABC 标题"]


class TestNormalizationOffsetIntegrity:
    """Offsets must remain valid AFTER normalization in registry.parse()."""

    def test_offsets_valid_after_normalization(self, tmp_path: Path) -> None:
        """Paragraphs with extra whitespace: offsets must match after normalize."""
        registry = create_default_registry()
        p = tmp_path / "spaces.txt"
        # Text with multiple spaces and trailing whitespace that normalize changes
        p.write_text("标题  行\n\n第一段   内容。  \n\n第二段\t\t内容。", encoding="utf-8")
        result = registry.parse(p)
        assert result.parse_status == "success"
        for para in result.paragraphs:
            extracted = result.raw_text[para.start_char : para.end_char]
            assert extracted == para.text, (
                f"Offset mismatch after normalization: "
                f"raw_text[{para.start_char}:{para.end_char}]="
                f"'{extracted}' != '{para.text}'"
            )

    def test_offsets_valid_with_fullwidth_letters(self, tmp_path: Path) -> None:
        """Full-width letters change char content but not length; offsets must hold."""
        registry = create_default_registry()
        p = tmp_path / "fw.txt"
        p.write_text("ＡＢＣ标题\n\nＤＥＦ内容。", encoding="utf-8")
        result = registry.parse(p)
        for para in result.paragraphs:
            extracted = result.raw_text[para.start_char : para.end_char]
            assert extracted == para.text

    def test_offsets_valid_md_after_normalization(self, tmp_path: Path) -> None:
        """Markdown through registry: offsets must survive normalization."""
        registry = create_default_registry()
        p = tmp_path / "test.md"
        p.write_text(
            "# 标题  行\n\n第一段   内容。\n\n## 第二节\n\n更多  内容。",
            encoding="utf-8",
        )
        result = registry.parse(p)
        assert result.parse_status == "success"
        for para in result.paragraphs:
            extracted = result.raw_text[para.start_char : para.end_char]
            assert extracted == para.text, (
                f"MD offset mismatch: raw_text[{para.start_char}:{para.end_char}]="
                f"'{extracted}' != '{para.text}'"
            )

    def test_offsets_valid_with_combining_chars(self, tmp_path: Path) -> None:
        """NFC normalization can change text length; offsets must still hold."""
        registry = create_default_registry()
        p = tmp_path / "nfc.txt"
        # e + combining acute → é (length changes from 2 to 1)
        p.write_text("caf\u0065\u0301\n\nre\u0301sume\u0301", encoding="utf-8")
        result = registry.parse(p)
        for para in result.paragraphs:
            extracted = result.raw_text[para.start_char : para.end_char]
            assert extracted == para.text


class TestParserChunkerIntegration:
    """End-to-end: parse a real file then chunk it, verify locators."""

    def test_txt_parse_then_chunk(self, tmp_path: Path) -> None:
        registry = create_default_registry()
        p = tmp_path / "doc.txt"
        p.write_text(
            "Title\n\n" + "A" * 500 + "\n\n" + "B" * 500,
            encoding="utf-8",
        )
        parsed = registry.parse(p)
        chunks = _chunk(parsed)
        assert len(chunks) >= 1
        for c in chunks:
            # char_start/char_end must locate valid text in raw_text
            assert 0 <= c.char_start <= c.char_end <= len(parsed.raw_text)

    def test_md_parse_then_chunk(self, tmp_md: Path) -> None:
        registry = create_default_registry()
        parsed = registry.parse(tmp_md)
        chunks = _chunk(parsed)
        assert len(chunks) >= 1
        # Heading paths should propagate to chunks
        headings = [c.heading_path for c in chunks if c.heading_path]
        assert len(headings) > 0
        # No raw markdown # prefix should leak into chunk text
        for c in chunks:
            for line in c.text.splitlines():
                stripped = line.strip()
                if stripped:
                    assert not stripped.startswith("# ")

    def test_md_chinese_parse_then_chunk(self, tmp_path: Path) -> None:
        registry = create_default_registry()
        p = tmp_path / "cn.md"
        p.write_text(
            "# 标题\n\n" + "中文内容。" * 200 + "\n\n## 第二节\n\n" + "更多内容。" * 200,
            encoding="utf-8",
        )
        parsed = registry.parse(p)
        chunks = _chunk(parsed)
        assert len(chunks) >= 2
        for c in chunks:
            assert 0 <= c.char_start <= c.char_end <= len(parsed.raw_text)

    def test_docx_parse_then_chunk(self, tmp_docx: Path) -> None:
        """Docx parse -> chunk: offsets and heading_path must be valid."""
        registry = create_default_registry()
        parsed = registry.parse(tmp_docx)
        assert parsed.parse_status == "success"
        chunks = _chunk(parsed)
        assert len(chunks) >= 1
        for c in chunks:
            assert 0 <= c.char_start <= c.char_end <= len(parsed.raw_text)
            located = parsed.raw_text[c.char_start : c.char_end]
            assert located in c.text
        # heading_path should propagate from docx styles
        headings = [c.heading_path for c in chunks if c.heading_path]
        assert len(headings) > 0

    def test_pdf_parse_then_chunk(self, tmp_pdf: Path) -> None:
        """PDF parse -> chunk: offsets and page_no must be valid."""
        registry = create_default_registry()
        parsed = registry.parse(tmp_pdf)
        assert parsed.parse_status == "success"
        chunks = _chunk(parsed)
        assert len(chunks) >= 1
        for c in chunks:
            assert 0 <= c.char_start <= c.char_end <= len(parsed.raw_text)
            located = parsed.raw_text[c.char_start : c.char_end]
            assert located in c.text
        # page_no should be preserved from PDF paragraphs
        page_nos = [c.page_no for c in chunks if c.page_no is not None]
        assert len(page_nos) > 0

    def test_multipage_pdf_chunk_page_no(self, tmp_pdf_multipage: Path) -> None:
        """Chunks from multi-page PDF should carry page_no from first paragraph."""
        registry = create_default_registry()
        parsed = registry.parse(tmp_pdf_multipage)
        assert parsed.parse_status == "success"
        chunks = _chunk(parsed)
        assert len(chunks) >= 1
        for c in chunks:
            assert c.page_no is not None
            assert c.page_no >= 1

    def test_pdf_toc_heading_path_in_chunks(self, tmp_pdf_with_toc: Path) -> None:
        """Chunks from PDF with TOC should carry heading_path."""
        registry = create_default_registry()
        parsed = registry.parse(tmp_pdf_with_toc)
        assert parsed.parse_status == "success"
        chunks = _chunk(parsed)
        assert len(chunks) >= 1
        paths = [c.heading_path for c in chunks if c.heading_path]
        assert len(paths) > 0, "Expected heading_path from TOC bookmarks"


class TestUltraShortDocuments:
    """Parse + chunk for documents with very few characters (< 10 chars)."""

    def test_ultra_short_txt(self, tmp_path: Path) -> None:
        registry = create_default_registry()
        p = tmp_path / "tiny.txt"
        p.write_text("OK", encoding="utf-8")
        parsed = registry.parse(p)
        assert parsed.parse_status == "success"
        chunks = _chunk(parsed)
        assert len(chunks) == 1
        assert chunks[0].text == "OK"

    def test_ultra_short_md(self, tmp_path: Path) -> None:
        registry = create_default_registry()
        p = tmp_path / "tiny.md"
        p.write_text("Hi", encoding="utf-8")
        parsed = registry.parse(p)
        assert parsed.parse_status == "success"
        chunks = _chunk(parsed)
        assert len(chunks) == 1
        assert "Hi" in chunks[0].text

    def test_ultra_short_docx(self, tmp_path: Path) -> None:
        pytest.importorskip("docx")
        from docx import Document  # type: ignore[import-untyped]

        registry = create_default_registry()
        doc = Document()
        doc.add_paragraph("AB")
        p = tmp_path / "tiny.docx"
        doc.save(str(p))
        parsed = registry.parse(p)
        assert parsed.parse_status == "success"
        chunks = _chunk(parsed)
        assert len(chunks) == 1
        assert "AB" in chunks[0].text

    def test_ultra_short_pdf(self, tmp_path: Path) -> None:
        fitz = pytest.importorskip("fitz")
        registry = create_default_registry()
        doc = fitz.open()
        page = doc.new_page()
        page.insert_text((72, 72), "XY")
        p = tmp_path / "tiny.pdf"
        doc.save(str(p))
        doc.close()
        parsed = registry.parse(p)
        assert parsed.parse_status == "success"
        chunks = _chunk(parsed)
        assert len(chunks) == 1
        assert "XY" in chunks[0].text


class TestParagraphValidation:
    """Paragraph __post_init__ must reject invalid field values."""

    def test_negative_start_char(self) -> None:
        with pytest.raises(ValueError, match="start_char"):
            Paragraph(text="x", index=0, start_char=-1, end_char=1)

    def test_end_char_before_start(self) -> None:
        with pytest.raises(ValueError, match="end_char"):
            Paragraph(text="x", index=0, start_char=5, end_char=3)

    def test_negative_index(self) -> None:
        with pytest.raises(ValueError, match="index"):
            Paragraph(text="x", index=-1, start_char=0, end_char=1)

    def test_page_no_zero(self) -> None:
        with pytest.raises(ValueError, match="page_no"):
            Paragraph(text="x", index=0, start_char=0, end_char=1, page_no=0)

    def test_valid_paragraph(self) -> None:
        p = Paragraph(text="hello", index=0, start_char=0, end_char=5, page_no=1)
        assert p.text == "hello"
        assert p.page_no == 1
