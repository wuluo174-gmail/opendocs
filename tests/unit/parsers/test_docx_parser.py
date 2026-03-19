"""Tests for DocxParser."""

from __future__ import annotations

from pathlib import Path

import pytest

docx = pytest.importorskip("docx")

from opendocs.parsers import docx_parser as docx_parser_module  # noqa: E402
from opendocs.parsers.docx_parser import DocxParser  # noqa: E402


class TestDocxParser:
    def setup_method(self) -> None:
        self.parser = DocxParser()

    def test_supported_extensions(self) -> None:
        assert self.parser.supported_extensions() == [".docx"]

    def test_parse_normal(self, tmp_docx: Path) -> None:
        result = self.parser.parse(tmp_docx)
        assert result.parse_status == "success"
        assert result.file_type == "docx"
        assert result.title == "Docx Title"
        assert len(result.paragraphs) >= 2

    def test_heading_extraction(self, tmp_docx: Path) -> None:
        result = self.parser.parse(tmp_docx)
        paths = [p.heading_path for p in result.paragraphs if p.heading_path]
        assert any("Heading One" in hp for hp in paths)
        assert any("Heading Two" in hp for hp in paths)

    def test_heading_hierarchy(self, tmp_docx: Path) -> None:
        result = self.parser.parse(tmp_docx)
        # Find paragraphs under Heading Two
        h2_paras = [
            p for p in result.paragraphs if p.heading_path and "Heading Two" in p.heading_path
        ]
        assert len(h2_paras) > 0
        # Heading Two (level 2) should be nested under Heading One (level 1)
        for p in h2_paras:
            assert "Heading One" in p.heading_path

    def test_char_offsets(self, tmp_docx: Path) -> None:
        result = self.parser.parse(tmp_docx)
        for para in result.paragraphs:
            assert para.start_char >= 0
            assert para.end_char >= para.start_char
            extracted = result.raw_text[para.start_char : para.end_char]
            assert extracted == para.text

    def test_chinese_content(self, tmp_docx_chinese: Path) -> None:
        result = self.parser.parse(tmp_docx_chinese)
        assert result.parse_status == "success"
        assert any("标题一" in p.text for p in result.paragraphs)
        assert any("第一段" in p.text for p in result.paragraphs)

    def test_empty_docx(self, tmp_docx_empty: Path) -> None:
        result = self.parser.parse(tmp_docx_empty)
        assert result.parse_status == "success"
        # Empty docx may still have some default paragraphs
        # but all should be empty-text filtered

    def test_corrupted_file(self, tmp_path: Path) -> None:
        """A corrupted .docx should raise ParseFailedError."""
        from opendocs.exceptions import ParseFailedError

        p = tmp_path / "bad.docx"
        p.write_bytes(b"this is not a docx")
        with pytest.raises(ParseFailedError):
            self.parser.parse(p)

    def test_char_offsets_match_raw_text(self, tmp_docx: Path) -> None:
        """Paragraph char offsets must align with raw_text content."""
        result = self.parser.parse(tmp_docx)
        for para in result.paragraphs:
            extracted = result.raw_text[para.start_char : para.end_char]
            assert extracted == para.text, (
                f"Offset mismatch: raw_text[{para.start_char}:{para.end_char}]="
                f"'{extracted}' != '{para.text}'"
            )

    def test_chinese_char_offsets(self, tmp_docx_chinese: Path) -> None:
        """Chinese docx paragraphs should have accurate char offsets."""
        result = self.parser.parse(tmp_docx_chinese)
        for para in result.paragraphs:
            extracted = result.raw_text[para.start_char : para.end_char]
            assert extracted == para.text

    def test_table_content_extracted(self, tmp_docx_with_table: Path) -> None:
        """Table cell text should be included in parsed paragraphs."""
        result = self.parser.parse(tmp_docx_with_table)
        all_text = result.raw_text
        assert "Alice" in all_text
        assert "95" in all_text
        assert "Name" in all_text

    def test_table_offsets_valid(self, tmp_docx_with_table: Path) -> None:
        """Table-derived paragraphs should have valid char offsets."""
        result = self.parser.parse(tmp_docx_with_table)
        for para in result.paragraphs:
            extracted = result.raw_text[para.start_char : para.end_char]
            assert extracted == para.text

    def test_hyperlink_text_not_lost(self, tmp_docx_with_hyperlink: Path) -> None:
        """Runs inside w:hyperlink containers must not be silently dropped."""
        result = self.parser.parse(tmp_docx_with_hyperlink)
        # The paragraph should contain the full text including the hyperlink run
        texts = [p.text for p in result.paragraphs]
        combined = " ".join(texts)
        assert "here" in combined, f"Hyperlink text lost; paragraphs: {texts}"
        # Specifically check the paragraph that has the link
        link_para = [p for p in result.paragraphs if "Click" in p.text]
        assert len(link_para) == 1
        assert "here" in link_para[0].text
        assert "please" in link_para[0].text

    def test_tabs_and_line_breaks_preserved(self, tmp_path: Path) -> None:
        """Inline tab and manual line break nodes must survive XML flattening."""
        from docx import Document  # type: ignore[import-untyped]
        from docx.enum.text import WD_BREAK  # type: ignore[import-untyped]

        doc = Document()
        para = doc.add_paragraph()
        run = para.add_run("Left")
        run.add_tab()
        run.add_text("Right")
        run = para.add_run()
        run.add_break(WD_BREAK.LINE)
        run.add_text("Next")

        p = tmp_path / "controls.docx"
        doc.save(str(p))

        result = self.parser.parse(p)
        assert result.parse_status == "success"
        assert len(result.paragraphs) == 1
        assert result.paragraphs[0].text == "Left\tRight\nNext"
        assert result.raw_text[result.paragraphs[0].start_char : result.paragraphs[0].end_char] == (
            "Left\tRight\nNext"
        )

    def test_partial_status_when_some_paragraphs_fail(self, tmp_path: Path) -> None:
        """S2-T04: DOCX with some unreadable paragraphs should yield partial."""
        from docx import Document  # type: ignore[import-untyped]

        doc = Document()
        doc.add_heading("Title", level=1)
        doc.add_paragraph("Good paragraph one.")
        doc.add_paragraph("Good paragraph two.")
        docx_path = tmp_path / "partial.docx"
        doc.save(str(docx_path))

        original_extract = docx_parser_module._extract_paragraph_text

        def _extract_with_single_failure(para_element, qn):
            text = original_extract(para_element, qn)
            if text == "Good paragraph two.":
                raise RuntimeError("broken paragraph")
            return text

        monkeypatch = pytest.MonkeyPatch()
        monkeypatch.setattr(
            docx_parser_module,
            "_extract_paragraph_text",
            _extract_with_single_failure,
        )
        try:
            result = self.parser.parse(docx_path)
        finally:
            monkeypatch.undo()

        assert result.parse_status == "partial"
        assert result.error_info is not None
        assert "2" in result.error_info
        assert result.error is not None
        assert result.error.code == "partial_parse"
        assert result.error.details["failed_source_paragraph_indices"] == [2]
        assert "source paragraphs" in result.error_info
        assert len(result.paragraphs) == 2
        assert result.raw_text == "Title\nGood paragraph one."
        for para in result.paragraphs:
            assert result.raw_text[para.start_char : para.end_char] == para.text

    def test_partial_status_when_table_block_fails(
        self,
        tmp_docx_with_table: Path,
        monkeypatch: pytest.MonkeyPatch,
    ) -> None:
        """S2-T04: broken table extraction must be auditable as partial."""
        import docx.table as docx_table  # type: ignore[import-untyped]

        def _broken_table(*_args, **_kwargs):
            raise RuntimeError("broken table")

        monkeypatch.setattr(docx_table, "Table", _broken_table)

        result = self.parser.parse(tmp_docx_with_table)

        assert result.parse_status == "partial"
        assert result.error_info is not None
        assert "failed table blocks" in result.error_info
        assert result.error is not None
        assert result.error.code == "partial_parse"
        assert result.error.details["parser"] == "DocxParser"
        assert "Intro text." in result.raw_text
        assert "Closing text." in result.raw_text
