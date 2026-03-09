"""Fixtures for parser tests – create temporary test files."""

from __future__ import annotations

from pathlib import Path

import pytest


@pytest.fixture()
def tmp_txt(tmp_path: Path) -> Path:
    """Create a simple .txt file with two paragraphs."""
    p = tmp_path / "sample.txt"
    p.write_text("Title Line\n\nFirst paragraph.\n\nSecond paragraph.", encoding="utf-8")
    return p


@pytest.fixture()
def tmp_txt_chinese(tmp_path: Path) -> Path:
    """Create a .txt file with Chinese content."""
    p = tmp_path / "chinese.txt"
    p.write_text("标题行\n\n第一段内容。\n\n第二段内容。", encoding="utf-8")
    return p


@pytest.fixture()
def tmp_txt_empty(tmp_path: Path) -> Path:
    """Create an empty .txt file."""
    p = tmp_path / "empty.txt"
    p.write_text("", encoding="utf-8")
    return p


@pytest.fixture()
def tmp_md(tmp_path: Path) -> Path:
    """Create a .md file with headings and paragraphs."""
    content = (
        "# Main Title\n"
        "\n"
        "Intro paragraph.\n"
        "\n"
        "## Section One\n"
        "\n"
        "Section one content.\n"
        "\n"
        "### Subsection A\n"
        "\n"
        "Subsection A content.\n"
        "\n"
        "## Section Two\n"
        "\n"
        "Section two content.\n"
    )
    p = tmp_path / "sample.md"
    p.write_text(content, encoding="utf-8")
    return p


@pytest.fixture()
def tmp_md_chinese(tmp_path: Path) -> Path:
    """Create a .md file with Chinese headings."""
    content = "# 引言\n\n背景说明。\n\n## 背景\n\n详细背景。\n"
    p = tmp_path / "chinese.md"
    p.write_text(content, encoding="utf-8")
    return p


@pytest.fixture()
def tmp_md_empty(tmp_path: Path) -> Path:
    p = tmp_path / "empty.md"
    p.write_text("", encoding="utf-8")
    return p


@pytest.fixture()
def tmp_docx(tmp_path: Path) -> Path:
    """Create a minimal .docx file with python-docx."""
    pytest.importorskip("docx")
    from docx import Document  # type: ignore[import-untyped]

    doc = Document()
    doc.core_properties.title = "Docx Title"
    doc.add_heading("Heading One", level=1)
    doc.add_paragraph("First paragraph.")
    doc.add_heading("Heading Two", level=2)
    doc.add_paragraph("Second paragraph.")
    p = tmp_path / "sample.docx"
    doc.save(str(p))
    return p


@pytest.fixture()
def tmp_docx_chinese(tmp_path: Path) -> Path:
    pytest.importorskip("docx")
    from docx import Document  # type: ignore[import-untyped]

    doc = Document()
    doc.add_heading("标题一", level=1)
    doc.add_paragraph("第一段。")
    p = tmp_path / "chinese.docx"
    doc.save(str(p))
    return p


@pytest.fixture()
def tmp_docx_with_table(tmp_path: Path) -> Path:
    """Create a .docx file containing a table."""
    pytest.importorskip("docx")
    from docx import Document  # type: ignore[import-untyped]

    doc = Document()
    doc.add_heading("Report", level=1)
    doc.add_paragraph("Intro text.")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Name"
    table.cell(0, 1).text = "Score"
    table.cell(1, 0).text = "Alice"
    table.cell(1, 1).text = "95"
    doc.add_paragraph("Closing text.")
    p = tmp_path / "table.docx"
    doc.save(str(p))
    return p


@pytest.fixture()
def tmp_md_frontmatter(tmp_path: Path) -> Path:
    """Create a .md file with YAML frontmatter."""
    content = (
        "---\n"
        "title: My Document\n"
        "date: 2026-01-01\n"
        "---\n"
        "\n"
        "# Actual Title\n"
        "\n"
        "Body content here.\n"
    )
    p = tmp_path / "frontmatter.md"
    p.write_text(content, encoding="utf-8")
    return p


@pytest.fixture()
def tmp_docx_empty(tmp_path: Path) -> Path:
    """Create a .docx with no paragraphs (just an empty doc)."""
    pytest.importorskip("docx")
    from docx import Document  # type: ignore[import-untyped]

    doc = Document()
    p = tmp_path / "empty.docx"
    doc.save(str(p))
    return p


@pytest.fixture()
def tmp_docx_with_hyperlink(tmp_path: Path) -> Path:
    """Create a .docx where a paragraph contains a hyperlink run."""
    pytest.importorskip("docx")
    from docx import Document  # type: ignore[import-untyped]
    from docx.oxml.ns import qn  # type: ignore[import-untyped]
    from docx.oxml import OxmlElement  # type: ignore[import-untyped]

    doc = Document()
    doc.add_heading("Title", level=1)
    para = doc.add_paragraph()
    # Add a normal run
    run1 = para.add_run("Click ")
    # Add a hyperlink element containing a run
    hyperlink = OxmlElement("w:hyperlink")
    run_in_link = OxmlElement("w:r")
    t_elem = OxmlElement("w:t")
    t_elem.text = "here"
    run_in_link.append(t_elem)
    hyperlink.append(run_in_link)
    para._element.append(hyperlink)
    # Add trailing normal run
    run2 = para.add_run(" please")

    p = tmp_path / "hyperlink.docx"
    doc.save(str(p))
    return p


@pytest.fixture()
def tmp_pdf(tmp_path: Path) -> Path:
    """Create a minimal PDF with PyMuPDF."""
    fitz = pytest.importorskip("fitz")
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 72), "PDF Title Line\n\nFirst paragraph.\n\nSecond paragraph.")
    p = tmp_path / "sample.pdf"
    doc.save(str(p))
    doc.close()
    return p


@pytest.fixture()
def tmp_pdf_chinese(tmp_path: Path) -> Path:
    fitz = pytest.importorskip("fitz")
    doc = fitz.open()
    page = doc.new_page()
    # PyMuPDF insert_text with CJK requires a CJK font; use insert_htmlbox as fallback
    try:
        page.insert_htmlbox(
            fitz.Rect(72, 72, 500, 700),
            "<p>标题行</p><p>第一段内容。</p>",
        )
    except Exception:
        page.insert_text((72, 72), "Chinese PDF")
    p = tmp_path / "chinese.pdf"
    doc.save(str(p))
    doc.close()
    return p


@pytest.fixture()
def tmp_pdf_empty(tmp_path: Path) -> Path:
    """Create a PDF with no text."""
    fitz = pytest.importorskip("fitz")
    doc = fitz.open()
    doc.new_page()  # blank page
    p = tmp_path / "empty.pdf"
    doc.save(str(p))
    doc.close()
    return p


@pytest.fixture()
def tmp_pdf_with_toc(tmp_path: Path) -> Path:
    """Create a multi-page PDF with TOC bookmarks for heading_path testing."""
    fitz = pytest.importorskip("fitz")
    doc = fitz.open()

    page1 = doc.new_page()
    page1.insert_text((72, 72), "Chapter 1 Introduction\n\nIntro content here.")

    page2 = doc.new_page()
    page2.insert_text((72, 72), "Chapter 2 Methods\n\nMethods content here.")

    page3 = doc.new_page()
    page3.insert_text((72, 72), "Section 2.1 Details\n\nDetailed methods.")

    # Add TOC bookmarks: [level, title, page_1based]
    doc.set_toc([
        [1, "Chapter 1 Introduction", 1],
        [1, "Chapter 2 Methods", 2],
        [2, "Section 2.1 Details", 3],
    ])

    p = tmp_path / "toc.pdf"
    doc.save(str(p))
    doc.close()
    return p


@pytest.fixture()
def tmp_pdf_multipage(tmp_path: Path) -> Path:
    """Create a multi-page PDF for cross-page chunk testing."""
    fitz = pytest.importorskip("fitz")
    doc = fitz.open()

    for i in range(1, 4):
        page = doc.new_page()
        page.insert_text((72, 72), f"Page {i} content paragraph.")

    p = tmp_path / "multipage.pdf"
    doc.save(str(p))
    doc.close()
    return p
